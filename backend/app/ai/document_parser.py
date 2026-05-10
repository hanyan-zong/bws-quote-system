"""文档解析 — PDF / DOCX / Excel / Image 分流处理后送 Claude."""
from __future__ import annotations

import base64
import json
import logging
from pathlib import Path
from typing import Any

from .claude_client import get_client

logger = logging.getLogger("bws.ai.parser")

SUPPORTED_EXTS = {".pdf", ".docx", ".xlsx", ".xls", ".png", ".jpg", ".jpeg", ".webp"}


SYSTEM_PROMPT = """你是一个专业的旅游成本数据抽取助手, 服务于一家面向中国市场的印尼地接旅行社。
你的工作是从供应商发来的报价文件中精准抽取以下类型的资源:
- hotel_room (酒店房型): hotel_name_zh, hotel_name_en, destination_code (DPS/CGK/SEPA/LOK/KMD), area, star, room_type, max_occupancy, breakfast_included, cost_idr_low, cost_idr_high, valid_from, valid_to, supplier, note
- attraction (景点): name_zh, destination_code, area, ticket_idr_adult, ticket_idr_child, recommended_minutes, note(可选)
- restaurant (餐厅): name_zh, cuisine, meal_type (lunch/dinner/both), area, cost_idr_per_person, min_pax, includes_drink
- vehicle (车辆): destination_code, seat_count, vehicle_type, cost_idr_per_day, includes_fuel, includes_driver, restrictions, note
- guide (导游): name_zh, destination_code, language (zh/en/id/zh+en), level (senior/regular/trainee), cost_idr_per_day, max_pax, note
- spa (SPA): brand, package_name, duration_minutes, cost_idr_per_person, note
- water_activity (水上项目): name_zh, location, cost_idr_per_person, duration_minutes
- afternoon_tea (下午茶): name_zh, venue, area, cost_idr_per_person
- optional_tour (自费): name_zh, sale_price_cny, cost_idr, target_audience, best_time, category

【价格清洗规则】(基于 BWS 母库实际数据)
- "Rp 25,000" / "Rp 25.000" / "Rp25,000" / "Rp25.000" → 25000(去千分位)
- "500k" / "500K" → 500000  /  "1.2M" → 1200000
- "fit 550k  7间房以上：480k" → cost_idr_low=480000, cost_idr_high=550000(取范围低值作 low)
- "650k-700k" → cost_idr_low=650000, cost_idr_high=700000
- "1.220.000" / "1,220,000" → 1220000
- "Rp 75，000"(中文逗号) → 75000
- 价格列出现 "/" 表示该档不可用,跳过(不要解析为 0)
- 单位若是"X 元/趟" 或 "X 元/人"(中文元) → 视为 CNY,标注 source_currency="CNY",并按 1 CNY = 2300 IDR 换算

【星级映射】
- "网评四星" / "国四" → star=4
- "网评五星" / "国五" / "国六" → star=5
- "网评三星" → star=3
- 不写星级 → null

【区域映射】(巴厘岛标准 31 个区域)
南部:努沙杜瓦/贝洛阿/库塔/勒吉安/水明漾/苍古/金巴兰/乌鲁瓦图/佩坎杜
中部:乌布/登巴萨/沙努尔/塔巴南
中北:百度库/金塔玛尼/巴杜尔
东部:卡朗加森/阿曼德/图兰本/齐齐党
北部:罗威纳/新加拉惹
西部:尼加拉/门吉里岛
外岛:佩尼达岛/蓝梦岛/切宁安岛/吉利岛/龙目岛/科莫多/瑟帕岛

【真实数据示例(few-shot)】
示例 1 — 酒店表:
  原文: "网评四星 | 巴厘岛百丽宫度假村酒店(Bali Paragon Resort Hotel) | fit 550k 7间房以上：480k"
  输出: {"resource_type":"hotel_room","data":{"hotel_name_zh":"巴厘岛百丽宫度假村酒店","hotel_name_en":"Bali Paragon Resort Hotel","star":4,"room_type":"Standard","cost_idr_low":480000,"cost_idr_high":550000,"max_occupancy":2,"breakfast_included":true,"note":"fit 550k 7间房以上:480k"}}

示例 2 — 门票表:
  原文: "1 | Waterblow 海底喷泉 | 门票 | Rp 25,000"
  输出: {"resource_type":"attraction","data":{"name_zh":"Waterblow 海底喷泉","ticket_idr_adult":25000,"ticket_idr_child":17500,"recommended_minutes":60,"area":"努沙杜瓦"}}
  说明: 儿童票通常是成人 70%(无明示时);Waterblow 在努沙杜瓦区(常识)

示例 3 — 车费(多档套餐):
  原文表头: 车型 | HALF DAY | FULL DAY | 4D/3N | 5D/4N | 6D/5N | 7D/6N
  原文: "TOYOTA HIACE | Rp 750.000 | Rp 1.050.000 | Rp 2.800.000 | ..."
  输出: {"resource_type":"vehicle","data":{"vehicle_type":"TOYOTA HIACE","seat_count":17,"cost_idr_per_day":1050000,"includes_fuel":true,"includes_driver":true,"note":"FULL DAY 报价;HALF DAY=750k;4D/3N 包价=2.8M"}}
  说明: cost_idr_per_day 取 FULL DAY;多日折扣价放 note

示例 4 — 导游小费规则:
  原文: "1-5人:160k/天/人,司机:65k/人/天"
  输出: {"resource_type":"guide","data":{"name_zh":"中文导游(团 1-5 人档)","language":"zh","level":"regular","cost_idr_per_day":160000,"max_pax":5,"note":"1-5 人团:160k 导游/65k 司机/天/人"}}

【输出规则】(严格遵守)
1. 仅返回 JSON, 不要任何解释性文字
2. 货币统一用 IDR;如果原文是 USD/CNY, 按 USD=15500 IDR / CNY=2300 IDR 换算并标注 source_currency
3. 不确定的字段填 null, 不要瞎猜
4. 一个文件可能含多个资源, 按数组返回
5. 给每条记录加 confidence (0~1), 低于 0.7 的字段在 low_confidence_fields 数组里列出
6. 中文+英文酒店名格式 "中文名(English Name)" → 拆成 hotel_name_zh + hotel_name_en

返回结构:
{
  "extraction_summary": "...",
  "resources": [
    {"resource_type": "hotel_room", "confidence": 0.95, "low_confidence_fields": [], "data": {...}}
  ],
  "warnings": ["..."]
}
"""


def detect_file_type(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix in (".xlsx", ".xls"):
        return "xlsx"
    if suffix in (".png", ".jpg", ".jpeg", ".webp"):
        return "image"
    raise ValueError(f"不支持的文件类型: {suffix}")


def extract_pdf_text(path: Path) -> str:
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        return f"[pdfplumber 未安装, 无法解析 {path.name}]"

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page_no, page in enumerate(pdf.pages, start=1):
            parts.append(f"\n===== Page {page_no} =====\n")
            tables = page.extract_tables()
            if tables:
                for ti, table in enumerate(tables, start=1):
                    parts.append(f"\n--- Table {ti} ---\n")
                    for row in table:
                        parts.append(" | ".join(str(c) if c else "" for c in row))
                        parts.append("\n")
            text = page.extract_text() or ""
            parts.append(text)
    return "".join(parts)


def extract_docx_text(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return f"[python-docx 未安装, 无法解析 {path.name}]"

    parts: list[str] = []
    doc = Document(str(path))
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for ti, table in enumerate(doc.tables, start=1):
        parts.append(f"\n--- Table {ti} ---")
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_xlsx(path: Path) -> dict[str, list[list[Any]]]:
    """支持 .xlsx (openpyxl) 和老格式 .xls (xlrd 1.2)."""
    suffix = path.suffix.lower()
    if suffix == ".xls":
        return _extract_xls_legacy(path)
    try:
        import openpyxl  # type: ignore
    except ImportError:
        return {"_error": [["openpyxl 未安装"]]}

    wb = openpyxl.load_workbook(str(path), data_only=True)
    out: dict[str, list[list[Any]]] = {}
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[list[Any]] = []
        for row in ws.iter_rows(values_only=True):
            if any(c is not None for c in row):
                rows.append([str(c) if c is not None else "" for c in row])
        out[sheet_name] = rows[:500]  # 限制行数避免 token 爆炸
    return out


def _extract_xls_legacy(path: Path) -> dict[str, list[list[Any]]]:
    """老格式 .xls 走 xlrd 1.2."""
    try:
        import xlrd  # type: ignore
    except ImportError:
        return {"_error": [["xlrd 未安装,无法解析 .xls。请把文件另存为 .xlsx 后重试"]]}
    try:
        wb = xlrd.open_workbook(str(path))
    except Exception as exc:
        return {"_error": [[f"xlrd 打开 .xls 失败: {exc}"]]}
    out: dict[str, list[list[Any]]] = {}
    for ws in wb.sheets():
        rows: list[list[Any]] = []
        for r in range(ws.nrows):
            vals = [ws.cell_value(r, c) for c in range(ws.ncols)]
            if any(v not in (None, "") for v in vals):
                rows.append([("" if v in (None, "") else str(v)) for v in vals])
        out[ws.name] = rows[:500]
    return out


TEMPLATE_SYSTEM_PROMPT = """你是一个专业的一日游行程模板抽取助手, 服务于巴厘岛地接旅行社。
从文档中识别一份"一日游"行程模板的结构化信息.

输出严格 JSON, 字段:
- name_zh: 行程中文名(如"乌布文化一日游")
- name_en: 英文名(可选)
- description: 简短描述(行程顺序)
- total_minutes_estimate: 全程估计分钟数(含驾驶+景点逗留+用餐)
- difficulty: easy/moderate/intense
- destination_code: DPS/CGK/SEPA/LOK/KMD
- attractions: [{name_zh, stay_minutes, order_index}] 按行程顺序
- restaurants: [{name_zh, meal_type: lunch/dinner/both}]
- warnings: 不确定/缺数据的说明

如果文档没明确说时长,合理推算.如果景点名是英文,翻译成中文.
仅返回 JSON 无任何额外文字."""


def parse_template_document(file_path: Path, hint: str | None = None) -> dict[str, Any]:
    """解析一日游模板文档 — 返回模板骨架 (name/desc/attractions/restaurants)."""
    file_type = detect_file_type(file_path)
    client = get_client()

    if client.mock_mode:
        result = client._mock_extract_template()  # type: ignore[attr-defined]
        result.setdefault("file_name", file_path.name)
        result.setdefault("file_type", file_type)
        return result

    # 走真实 Claude 时复用文本抽取逻辑(图片暂不接,模板通常是文档)
    if file_type == "pdf":
        text = extract_pdf_text(file_path)
    elif file_type == "docx":
        text = extract_docx_text(file_path)
    elif file_type == "xlsx":
        sheets = extract_xlsx(file_path)
        text = json.dumps(sheets, ensure_ascii=False, indent=2)
    else:
        return {"_error": "模板暂不支持图片解析", "warnings": ["请用 Excel/DOCX/PDF"]}

    blocks = [{"type": "text", "text": f"以下是一日游模板文件 {file_path.name}.提示: {hint or '无'}\n\n{text}"}]
    result = client.extract_json(system=TEMPLATE_SYSTEM_PROMPT, content_blocks=blocks)
    result.setdefault("file_name", file_path.name)
    result.setdefault("file_type", file_type)
    return result


ITINERARY_SYSTEM_PROMPT = """你是一个专业的旅游行程意向解析助手, 服务于一家面向中国市场的印尼地接旅行社的 B 端业务员。
B 端旅行社 (即客户) 把"我想要的行程"以 PDF / Word / Excel / 图片形式上传, 你要从中精准抽取出**报价意向**, 让业务员能一键生成报价。

【需要抽取的字段】
基本信息:
- agency_name: B 端旅行社名 (可选)
- customer_name: 客户/团队名称 (可选)
- pax_adult: 成人数 (>= 1)
- pax_child: 儿童数 (默认 0; 文件中"小孩 X 人"/"儿童 X 名")
- pax_senior: 55 岁以上老年人数 (默认 0; 文件中"老人 X 人"/"55+"等)
- start_date: 出发日期 (YYYY-MM-DD)
- end_date: 结束日期
- total_days: 总天数 (= end_date - start_date + 1)
- destination_codes: 目的地代码数组, 仅用 [DPS=巴厘岛, CGK=雅加达, SEPA=瑟帕岛, LOK=龙目岛, KMD=科莫多]
- season: 自动判断 (1-3,11-12=high; 4,5,9,10=shoulder; 6,7,8=low) — 6/7/8 因雨季在巴厘岛是 low
- customer_type: 从 [honeymoon=蜜月, family_kids=亲子, young=年轻人, family=家庭, senior=长辈, mice=MICE/会奖, wedding=婚礼] 推断
- is_first_time_agency: 默认 false
- arrival_at / departure_at: 抵达/离开航班时间 (ISO datetime, 可选)
- arrival_airport / departure_airport: 三字代码 (可选, 如 DPS/CGK)
- notes: 客户特殊要求文字汇总

每日行程 days[]:
- day_index: 第几天 (从 1 开始)
- date: YYYY-MM-DD
- is_free: 是否全天自由 (true/false)
- free_hours: 自由时间小时 (0-12)
- hotel_name: AI 抽到的酒店原文名 (后端做匹配, 你不需要给 ID)
- hotel_star_request: 客户想要的酒店星级 (3/4/5; 写"五星"=5)
- room_type_request: 房型要求 (如 "海景房" "含早")
- vehicle_request: 车辆要求 ("商务车" "7 座 SUV" "中巴")
- guide_required: 是否要导游
- breakfast_included: 含早?
- lunch_request / dinner_request: 餐厅意向
- attractions: 景点名字数组 (如 ["乌布皇宫", "圣猴森林"])
- spa_request / water_activity_request / afternoon_tea_request: 升级体验
- notes: 当日特殊要求

【返回 JSON 格式】
{
  "extraction_summary": "解析了一份蜜月行程, 共 5 天 4 晚, 2 大人 1 小孩, 含瑟帕岛跳岛",
  "confidence": 0.85,
  "quote_draft": {
    "agency_name": "...", "customer_name": "...",
    "pax_adult": 2, "pax_child": 1, "pax_senior": 0,
    "start_date": "2026-06-01", "end_date": "2026-06-05",
    "total_days": 5, "destination_codes": ["DPS","SEPA"],
    "season": "high", "customer_type": "honeymoon",
    "notes": "客户要求海景房+私密浪漫",
    "days": [
      {"day_index": 1, "date": "2026-06-01", "is_free": false, "free_hours": 0,
       "hotel_name": "巴厘岛 The Mulia", "hotel_star_request": 5, "room_type_request": "海景房",
       "vehicle_request": "商务车", "breakfast_included": true,
       "lunch_request": "脏鸭餐厅", "attractions": ["乌布皇宫","圣猴森林"],
       "notes": ""},
      ...
    ]
  },
  "warnings": ["未明确指定 Day 3 酒店, 请补充"],
  "missing_fields": ["pax_senior", "days[2].hotel"]
}

【重要规则】
- 不确定的字段返回 null, 不要瞎猜
- 缺失但关键的字段 (人数/日期/目的地/天数), 务必加到 missing_fields
- pax_adult 必须 >= 1, 如果完全找不到就填 missing_fields
- 日期范围或单一日期都接受 (单日填 start_date=end_date, total_days=1)
"""


def parse_itinerary_intent(file_path: Path, hint: str | None = None) -> dict[str, Any]:
    """v0.6 — 解析客户上传的"行程意向"文件 (与 parse_document 抽资源不同).

    返回 quote_draft + missing_fields, 供前端做"补漏 → 一键报价"流程.
    """
    file_type = detect_file_type(file_path)
    client = get_client()

    if client.mock_mode:
        # 返回一个 mock 行程意向, 让前端流程能跑通
        return {
            "extraction_summary": "[MOCK] 蜜月 5 天 4 晚, 2 大 0 小, 巴厘岛, 含瑟帕岛跳岛",
            "confidence": 0.8,
            "quote_draft": {
                "agency_name": "(请填写 B 端旅行社名)",
                "customer_name": "(从文件名/上下文推断)",
                "pax_adult": 2, "pax_child": 0, "pax_senior": 0,
                "start_date": None, "end_date": None, "total_days": 5,
                "destination_codes": ["DPS"],
                "season": "shoulder", "customer_type": "honeymoon",
                "notes": "[MOCK] AI 模拟数据 — 未配置 ANTHROPIC_API_KEY",
                "days": [
                    {"day_index": i, "date": None, "is_free": (i in (3, 4)), "free_hours": (8 if i in (3, 4) else 0),
                     "hotel_name": "巴厘岛 The Mulia", "hotel_star_request": 5, "room_type_request": "海景房",
                     "vehicle_request": "商务车", "breakfast_included": True,
                     "lunch_request": None, "dinner_request": None,
                     "attractions": ["乌布皇宫" if i == 1 else "金巴兰沙滩"], "notes": ""}
                    for i in range(1, 6)
                ],
            },
            "warnings": ["MOCK 数据 — 真实 AI 解析需 ANTHROPIC_API_KEY"],
            "missing_fields": ["start_date", "end_date", "agency_name", "customer_name"],
            "file_name": file_path.name,
            "file_type": file_type,
        }

    if file_type == "pdf":
        text = extract_pdf_text(file_path)
        content_blocks = [{"type": "text", "text": f"以下是 PDF {file_path.name} 内容(行程意向). 提示: {hint or '无'}\n\n{text}"}]
    elif file_type == "docx":
        text = extract_docx_text(file_path)
        content_blocks = [{"type": "text", "text": f"以下是 DOCX {file_path.name} 内容(行程意向). 提示: {hint or '无'}\n\n{text}"}]
    elif file_type == "xlsx":
        sheets = extract_xlsx(file_path)
        content_blocks = [{"type": "text", "text": f"以下是 Excel {file_path.name} 内容(行程意向). 提示: {hint or '无'}\n\n" + json.dumps(sheets, ensure_ascii=False, indent=2)}]
    else:
        try:
            data = file_path.read_bytes()
            b64 = base64.standard_b64encode(data).decode()
            ext = file_path.suffix.lower().lstrip(".")
            media_type = "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"
            content_blocks = [
                {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64}},
                {"type": "text", "text": f"上图是客户行程意向. 提示: {hint or '无'}"},
            ]
        except Exception as exc:
            return {"_error": str(exc), "quote_draft": None, "warnings": [str(exc)],
                    "missing_fields": ["everything"], "file_name": file_path.name, "file_type": file_type}

    result = client.extract_json(system=ITINERARY_SYSTEM_PROMPT, content_blocks=content_blocks)
    result.setdefault("file_name", file_path.name)
    result.setdefault("file_type", file_type)
    result.setdefault("missing_fields", [])
    result.setdefault("warnings", [])
    return result


def parse_document(file_path: Path, hint: str | None = None) -> dict[str, Any]:
    """主入口 — 给定文件路径返回 AI 解析的结构化结果."""
    file_type = detect_file_type(file_path)
    client = get_client()

    # mock 模式下跳过本地解析(否则假文件会让 pdfplumber/openpyxl 炸)
    if client.mock_mode:
        result = client._mock_extract([])
        result.setdefault("file_name", file_path.name)
        result.setdefault("file_type", file_type)
        return result

    if file_type == "pdf":
        text = extract_pdf_text(file_path)
        content_blocks = [
            {"type": "text", "text": f"以下是 PDF 文件 {file_path.name} 的内容 (含表格)。提示: {hint or '无'}\n\n{text}"}
        ]
    elif file_type == "docx":
        text = extract_docx_text(file_path)
        content_blocks = [
            {"type": "text", "text": f"以下是 DOCX 文件 {file_path.name} 的内容。提示: {hint or '无'}\n\n{text}"}
        ]
    elif file_type == "xlsx":
        sheets = extract_xlsx(file_path)
        content_blocks = [
            {
                "type": "text",
                "text": f"以下是 Excel 文件 {file_path.name} 各 sheet 的内容 (JSON 格式)。提示: {hint or '无'}\n\n"
                + json.dumps(sheets, ensure_ascii=False, indent=2),
            }
        ]
    else:  # image
        try:
            data = file_path.read_bytes()
            b64 = base64.standard_b64encode(data).decode()
            ext = file_path.suffix.lower().lstrip(".")
            media_type = "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"
            content_blocks = [
                {
                    "type": "image",
                    "source": {"type": "base64", "media_type": media_type, "data": b64},
                },
                {"type": "text", "text": f"上图是供应商报价单截图。提示: {hint or '无'}"},
            ]
        except Exception as exc:
            logger.exception("读取图片失败")
            return {"_error": str(exc), "resources": [], "warnings": [str(exc)]}

    result = client.extract_json(system=SYSTEM_PROMPT, content_blocks=content_blocks)
    result.setdefault("file_name", file_path.name)
    result.setdefault("file_type", file_type)
    return result
