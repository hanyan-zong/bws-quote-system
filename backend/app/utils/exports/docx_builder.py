"""报价单 Word 导出 — 复用母库雅加达瑟帕岛报价排版风格."""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# BWS 品牌色
BRAND_BLUE = RGBColor(0x2C, 0x52, 0x82)
BRAND_LIGHT_HEX = "EBF8FF"
GRAY_TEXT = RGBColor(0x4A, 0x55, 0x68)


def _set_font(run, *, size: int = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    # 中文字体兼容
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _shade_cell(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_font(run, size=14, bold=True, color=BRAND_BLUE)


def _add_kv_table(doc: Document, rows: list[tuple[str, str, str, str]]) -> None:
    """4 列表: label1 | val1 | label2 | val2"""
    t = doc.add_table(rows=len(rows), cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(3), Cm(5), Cm(3), Cm(5)]
    for i, (l1, v1, l2, v2) in enumerate(rows):
        cells = t.rows[i].cells
        for j, (txt, is_label) in enumerate([(l1, True), (v1, False), (l2, True), (v2, False)]):
            cell = cells[j]
            cell.width = widths[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            run = cell.paragraphs[0].add_run(txt)
            if is_label:
                _set_font(run, bold=True, color=BRAND_BLUE)
                _shade_cell(cell, BRAND_LIGHT_HEX)
            else:
                _set_font(run)


def build_docx(ctx: dict[str, Any]) -> bytes:
    """生成 Word 报价单."""
    q = ctx["quote"]
    show = ctx["show_costs"]

    doc = Document()
    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ---- 标题 ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BWS 预报价单")
    _set_font(run, size=22, bold=True, color=BRAND_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"报价单号 {q['quote_no']}    |    导出时间 {ctx['meta']['exported_at']}    |    导出 {ctx['meta']['exported_by']}"
    )
    _set_font(run, size=9, color=GRAY_TEXT)

    # ---- 一、客户与行程概要 ----
    _add_heading(doc, "一、客户与行程概要")
    info_rows = [
        ("旅行社", q["agency_name"] or "—", "联系人", q["agency_contact"] or "—"),
        ("客户姓名", q["customer_name"] or "—", "目的地", q["destination_codes"] or "—"),
        ("出发 / 结束", f"{q['start_date']} → {q['end_date']}", "天数 / 自由", f"{q['total_days']} 天 / {q['free_days']} 自由日"),
        ("成人 / 儿童", f"{q['pax_adult']} 大 + {q['pax_child']} 小 = {q['pax_total']} 人",
         "客户类型 / 季节", f"{q['customer_type_label']} / {q['season_label']}"),
    ]
    if q["arrival_at"] or q["departure_at"]:
        info_rows.append((
            "抵达航班", f"{q['arrival_at']} {q['arrival_airport']}",
            "离开航班", f"{q['departure_at']} {q['departure_airport']}",
        ))
    _add_kv_table(doc, info_rows)

    # ---- 二、逐日行程 ----
    _add_heading(doc, "二、逐日行程")
    for d in ctx["days"]:
        # 标题段
        p = doc.add_paragraph()
        run = p.add_run(f"第 {d['day_index']} 天    {d['date']}")
        _set_font(run, size=12, bold=True, color=BRAND_BLUE)
        if d["is_free"]:
            tag = p.add_run(f"    [全天自由 {d['free_hours']}h]")
            _set_font(tag, size=9, color=GRAY_TEXT)
        elif d["free_hours"] >= 4:
            tag = p.add_run(f"    [半自由 {d['free_hours']}h]")
            _set_font(tag, size=9, color=GRAY_TEXT)

        kv: list[tuple[str, str]] = []
        if d["hotel"]:
            kv.append(("住宿", f"{d['hotel']}{(' / ' + d['room']) if d['room'] else ''}"))
        if d["vehicle"] or d["guide"]:
            kv.append(("用车 / 导游", f"{d['vehicle']}{(' | 导游 ' + d['guide']) if d['guide'] else ''}"))
        meals_parts = []
        if d["breakfast_included"]:
            meals_parts.append("含早")
        if d["lunch"]:
            meals_parts.append(f"午: {d['lunch']}")
        if d["dinner"]:
            meals_parts.append(f"晚: {d['dinner']}")
        if d["afternoon_tea"]:
            meals_parts.append(f"下午茶: {d['afternoon_tea']}")
        if meals_parts:
            kv.append(("餐", " · ".join(meals_parts)))
        if d["attractions"]:
            attr_lines = "\n".join(
                f"  {a['order']}. {a['name']}{(' (' + str(a['stay_minutes']) + ' min)') if a['stay_minutes'] else ''}"
                for a in d["attractions"]
            )
            kv.append(("景点 / 体验", attr_lines))
        if d["spa"] or d["water_activity"]:
            up = []
            if d["spa"]:
                up.append(f"SPA: {d['spa']}")
            if d["water_activity"]:
                up.append(f"水上: {d['water_activity']}")
            kv.append(("升级体验", " | ".join(up)))
        if d["notes"]:
            kv.append(("备注", d["notes"]))

        if kv:
            t = doc.add_table(rows=len(kv), cols=2)
            t.style = "Table Grid"
            for i, (label, val) in enumerate(kv):
                cells = t.rows[i].cells
                cells[0].width = Cm(3)
                cells[1].width = Cm(13)
                cells[0].text = ""
                run = cells[0].paragraphs[0].add_run(label)
                _set_font(run, bold=True, color=BRAND_BLUE)
                _shade_cell(cells[0], BRAND_LIGHT_HEX)
                cells[1].text = ""
                # 多行内容
                for line_i, line in enumerate(val.split("\n")):
                    if line_i == 0:
                        para = cells[1].paragraphs[0]
                    else:
                        para = cells[1].add_paragraph()
                    run = para.add_run(line)
                    _set_font(run)
        doc.add_paragraph()  # 分隔

    # ---- 三、价格汇总 ----
    _add_heading(doc, "三、价格汇总")
    price_rows: list[tuple[str, str]] = [
        ("人均报价", f"¥ {q['price_cny_per_pax']:,.2f}"),
        ("总报价", f"¥ {q['price_cny_total']:,.2f}"),
    ]
    if show:
        price_rows.extend([
            ("人均成本 (CNY)", f"¥ {q.get('cost_cny_total', 0) / q['pax_total']:,.2f}"),
            ("人均利润", f"¥ {q.get('profit_cny_per_pax', 0):,.2f}"),
            ("人均赌额 (让利)", f"¥ {q.get('gamble_cny_per_pax', 0):,.2f}"),
            ("汇率 RMB:IDR", f"1 : {q['exchange_rate']:,.2f}"),
        ])
    t = doc.add_table(rows=len(price_rows), cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, val) in enumerate(price_rows):
        cells = t.rows[i].cells
        cells[0].width = Cm(6)
        cells[1].width = Cm(10)
        cells[0].text = ""
        cells[1].text = ""
        l_run = cells[0].paragraphs[0].add_run(label)
        _set_font(l_run, bold=True, color=BRAND_BLUE)
        _shade_cell(cells[0], BRAND_LIGHT_HEX)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        v_run = cells[1].paragraphs[0].add_run(val)
        _set_font(v_run, bold=True)

    if not show:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("※ 本报价为最终客户成交价 · 内部成本与利润数据已隐藏")
        _set_font(run, size=9, color=GRAY_TEXT)
        run.italic = True

    if q["notes"]:
        _add_heading(doc, "四、备注")
        p = doc.add_paragraph()
        run = p.add_run(q["notes"])
        _set_font(run)

    # 页脚
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"本报价单由 BWS 预报价系统 B 端版本 (v{ctx['meta']['version']}) 自动生成 · © PT BWS Indonesia"
    )
    _set_font(run, size=8, color=GRAY_TEXT)
    run.italic = True

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
